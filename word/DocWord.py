import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING

document = Document()


class DocWord:
    # default constructor
    def __init__(self, details):
        self.add_name(details.name)
        self.add_phone_and_email(details.phone_num, details.email)
        document.add_heading('Work Experience', level=1)
        document.save('resume_demo.docx')

    def add_name(self, name):
        paragraph_name = document.add_paragraph(name)
        document.save('resume_demo.docx')

        style = document.styles['Normal']
        font = style.font
        font.name = 'Times new roman'
        font.size = Pt(11)
        paragraph_name.style = document.styles['Normal']
        paragraph_name.style.font.bold = 1
        paragraph_name.alignment = 1
        paragraph_name.paragraph_format.space_before = Pt(0)
        paragraph_name.paragraph_format.space_after = Pt(0)
        document.save('resume_demo.docx')



    def create_new_word_file(self):
        #document.add_heading('Document Title', 0)
        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True
        document.add_heading('Heading, level 1', level=1)
        #document.add_paragraph('Intense quote', style='Intense Quote')
        self.add_bullet_to_resume("aabbcc")
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )
        document.add_picture("eiffel-tower.jpg", width=Inches(1.25))
        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc
        document.add_page_break()
        document.save('demo.docx')
        return document

    def add_bullet_to_resume(self, bullets):
        for bullet in bullets:
            document.add_paragraph(
                bullet[0], style='List Bullet'
            )
        document.save('resume_demo.docx')

    # def word():
    #     mydoc = docx.Document()
    #     mydoc.add_paragraph("This is first paragraph of a MS Word file.")
    #     mydoc = docx.Document("resume.docx")
    #     mydoc.add_paragraph("This is the second paragraph of a MS Word file.")
    #     mydoc = docx.Document("resume.docx")
    #     third_para = mydoc.add_paragraph("This is the third paragraph.")
    #     third_para.add_run(" this is a section at the end of third paragraph")
    #     mydoc.save("resume_template.docx")
    #     mydoc.add_heading("This is level 1 heading", 0)
    #     mydoc.add_heading("This is level 2 heading", 1)
    #     mydoc.add_heading("This is level 3 heading", 2)
    #     mydoc.save("resume_template.docx")
    #     mydoc.add_picture("eiffel-tower.jpg", width=docx.shared.Inches(2), height=docx.shared.Inches(3))
    #     mydoc.save("resume_template.docx")


    def add_phone_and_email(self, phone_num, email):
        if phone_num.__len__() == 10:
            phone_num = phone_num[: 0] + "(" + phone_num[0:]
            phone_num = phone_num[: 4] + ")" + phone_num[4:]
            phone_num = phone_num[: 8] + "-" + phone_num[8:]
        phone_and_email = phone_num + " | " + email
        paragraph = document.add_paragraph(phone_and_email)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style = document.styles['Normal']
        font = style.font
        font.name = 'Times new roman'
        font.size = Pt(10)
        paragraph.style = document.styles['Normal']
        paragraph.style.font.bold = 0
        paragraph.alignment = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1)
        document.save('resume_demo.docx')





class details_test:
    name = "ANTHONY (TONY) STARK"
    phone_num = "8572851027"
    email = "assafy@mit.edu"

if __name__ == '__main__':
    obj = DocWord(details_test)
    list1 = ["Aaa","bbb","11111"]
    obj.add_bullet_to_resume(list1)

